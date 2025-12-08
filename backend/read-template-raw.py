"""
Script para ver o texto RAW do documento e detectar formato das variáveis
"""

from pathlib import Path
from docx import Document

template_path = r"C:\Users\roger\Downloads\Proposta 2025 - CORRIGIDO.docx"

if not Path(template_path).exists():
    print(f"❌ Arquivo não encontrado")
    exit(1)

doc = Document(template_path)

print("\n" + "="*80)
print("📄 TEXTO DOS PRIMEIROS 20 PARÁGRAFOS")
print("="*80 + "\n")

for i, para in enumerate(doc.paragraphs[:20]):
    if para.text.strip():
        print(f"[{i}] {para.text}")

print("\n" + "="*80)
print("📊 TEXTO DAS PRIMEIRAS CÉLULAS DAS TABELAS")
print("="*80 + "\n")

for t_idx, table in enumerate(doc.tables[:3]):
    print(f"\n--- TABELA {t_idx+1} ---")
    for r_idx, row in enumerate(table.rows[:5]):
        for c_idx, cell in enumerate(row.cells[:5]):
            text = cell.text.strip()
            if text:
                print(f"  [{r_idx},{c_idx}] {text}")
