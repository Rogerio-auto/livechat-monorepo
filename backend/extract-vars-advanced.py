"""
Extrai variáveis de campos do Word (fields) e content controls
"""

from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
import re

template_path = r"C:\Users\roger\Downloads\Proposta 2025 - CORRIGIDO.docx"

if not Path(template_path).exists():
    print(f"❌ Arquivo não encontrado")
    exit(1)

doc = Document(template_path)

print("\n" + "="*80)
print("🔍 ANALISANDO ESTRUTURA XML DO DOCUMENTO")
print("="*80 + "\n")

variables = set()

# Função para extrair texto de runs (incluindo fields)
def extract_text_from_element(element):
    """Extrai texto incluindo de fields"""
    texts = []
    
    # Texto direto
    for t_elem in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        if t_elem.text:
            texts.append(t_elem.text)
    
    # Fields (fldSimple)
    for fld in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldSimple'):
        instr = fld.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr', '')
        if instr:
            texts.append(f"FIELD:{instr}")
    
    # Fields complexos (fldChar)
    return ' '.join(texts)

# 1. Parágrafos
print("📝 Parágrafos com conteúdo:")
for i, para in enumerate(doc.paragraphs):
    text = extract_text_from_element(para._element)
    if text.strip():
        print(f"  [{i}] {text[:200]}")
        # Procurar padrões de variáveis
        vars_found = re.findall(r'\{\{\s*\w+[\w.]*\s*\}\}', text)
        if vars_found:
            print(f"       Variáveis: {vars_found}")
            variables.update(vars_found)

# 2. Tabelas
print(f"\n📊 Tabelas ({len(doc.tables)} encontradas):")
for t_idx, table in enumerate(doc.tables):
    print(f"\n  TABELA {t_idx + 1}:")
    for r_idx, row in enumerate(table.rows[:3]):  # Primeiras 3 linhas
        for c_idx, cell in enumerate(row.cells[:5]):  # Primeiras 5 colunas
            text = extract_text_from_element(cell._element)
            if text.strip():
                print(f"    [{r_idx},{c_idx}] {text[:100]}")
                vars_found = re.findall(r'\{\{\s*\w+[\w.]*\s*\}\}', text)
                if vars_found:
                    print(f"             Variáveis: {vars_found}")
                    variables.update(vars_found)

# 3. Content Controls
print("\n🎛️ Content Controls:")
cc_elements = doc.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt')
print(f"  Encontrados: {len(cc_elements)}")

for i, cc in enumerate(cc_elements[:10]):  # Primeiros 10
    text = extract_text_from_element(cc)
    if text.strip():
        print(f"  [{i}] {text[:150]}")

print(f"\n{'='*80}")
print(f"📋 VARIÁVEIS ENCONTRADAS: {len(variables)}")
print(f"{'='*80}\n")

for var in sorted(variables):
    print(f"  {var}")
