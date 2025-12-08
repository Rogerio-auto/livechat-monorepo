"""
Script para extrair TODAS as variáveis de um documento DOCX
Identifica variáveis simples {{ var }} e loops {% for %}
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("ERROR: python-docx não instalado")
    print("Execute: pip install python-docx")
    sys.exit(1)


def extract_variables_from_text(text):
    """Extrai variáveis {{ var }} e loops {% for %} do texto"""
    variables = set()
    loops = []
    
    # Variáveis simples: {{ NOME }}
    simple_vars = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
    for var in simple_vars:
        var = var.strip()
        # Remover filtros (ex: {{ valor|default:"" }})
        var = var.split('|')[0].strip()
        variables.add(var)
    
    # Loops: {% for item in items %}
    for_loops = re.findall(r'\{%\s*for\s+(\w+)\s+in\s+(\w+)\s*%\}', text)
    for item_var, collection_var in for_loops:
        loops.append({
            'collection': collection_var,
            'item': item_var,
            'variables': set()
        })
    
    # Variáveis dentro de loops: {{ item.campo }}
    loop_vars = re.findall(r'\{\{\s*(\w+)\.(\w+)\s*\}\}', text)
    for obj, field in loop_vars:
        # Adicionar à variável composta
        variables.add(f"{obj}.{field}")
        
        # Se é uma variável de loop, adicionar ao loop correspondente
        for loop in loops:
            if obj == loop['item']:
                loop['variables'].add(field)
    
    return variables, loops


def extract_all_variables(docx_path):
    """Extrai todas as variáveis do documento"""
    print(f"\n{'='*80}")
    print(f"📄 ANALISANDO: {docx_path}")
    print(f"{'='*80}\n")
    
    doc = Document(docx_path)
    
    all_variables = set()
    all_loops = []
    
    # 1. Parágrafos
    print("🔍 Analisando parágrafos...")
    for para in doc.paragraphs:
        if para.text:
            vars, loops = extract_variables_from_text(para.text)
            all_variables.update(vars)
            all_loops.extend(loops)
    
    # 2. Tabelas
    print("🔍 Analisando tabelas...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        vars, loops = extract_variables_from_text(para.text)
                        all_variables.update(vars)
                        all_loops.extend(loops)
    
    # 3. Cabeçalhos e rodapés
    print("🔍 Analisando cabeçalhos e rodapés...")
    for section in doc.sections:
        # Cabeçalho
        if section.header:
            for para in section.header.paragraphs:
                if para.text:
                    vars, loops = extract_variables_from_text(para.text)
                    all_variables.update(vars)
                    all_loops.extend(loops)
        
        # Rodapé
        if section.footer:
            for para in section.footer.paragraphs:
                if para.text:
                    vars, loops = extract_variables_from_text(para.text)
                    all_variables.update(vars)
                    all_loops.extend(loops)
    
    # Separar variáveis simples de variáveis de loops
    simple_vars = {v for v in all_variables if '.' not in v}
    loop_vars = {v for v in all_variables if '.' in v}
    
    return simple_vars, loop_vars, all_loops


def print_results(simple_vars, loop_vars, all_loops):
    """Imprime resultados formatados"""
    
    print(f"\n{'='*80}")
    print("📊 VARIÁVEIS SIMPLES")
    print(f"{'='*80}\n")
    
    if simple_vars:
        for var in sorted(simple_vars):
            print(f"  {{ {var} }}")
    else:
        print("  (nenhuma encontrada)")
    
    print(f"\n{'='*80}")
    print("🔁 LOOPS E SUAS VARIÁVEIS")
    print(f"{'='*80}\n")
    
    # Consolidar loops por collection
    loops_by_collection = {}
    for loop in all_loops:
        coll = loop['collection']
        if coll not in loops_by_collection:
            loops_by_collection[coll] = {
                'item_name': loop['item'],
                'variables': set()
            }
        loops_by_collection[coll]['variables'].update(loop['variables'])
    
    if loops_by_collection:
        for collection, data in sorted(loops_by_collection.items()):
            print(f"  {{% for {data['item_name']} in {collection} %}}")
            if data['variables']:
                for var in sorted(data['variables']):
                    print(f"    {{ {data['item_name']}.{var} }}")
            else:
                print("    (nenhuma variável detectada)")
            print(f"  {{% endfor %}}\n")
    else:
        print("  (nenhum encontrado)")
    
    print(f"\n{'='*80}")
    print("📋 RESUMO")
    print(f"{'='*80}\n")
    print(f"  Variáveis simples: {len(simple_vars)}")
    print(f"  Variáveis de loops: {len(loop_vars)}")
    print(f"  Loops encontrados: {len(loops_by_collection)}")
    print(f"\n{'='*80}\n")


if __name__ == "__main__":
    # Caminho do template
    template_path = r"C:\Users\roger\Downloads\Proposta 2025 - CORRIGIDO.docx"
    
    if not Path(template_path).exists():
        print(f"❌ Arquivo não encontrado: {template_path}")
        sys.exit(1)
    
    # Extrair variáveis
    simple_vars, loop_vars, all_loops = extract_all_variables(template_path)
    
    # Imprimir resultados
    print_results(simple_vars, loop_vars, all_loops)
